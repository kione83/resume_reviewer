#!/usr/bin/env python3
"""
PMP resume scoring folder watcher.

When a PDF or DOCX resume lands in WATCH_FOLDER, this script extracts the
resume text, asks OpenAI for strict JSON rubric scores, appends only raw scores
and rationale to Google Sheets, then moves the file to processed/ or failed/.
"""

from __future__ import annotations

import hashlib
import argparse
import base64
import io
import json
import logging
import os
import random
import re
import shutil
import subprocess
import sys
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import gspread
import pdfplumber
import pypdfium2 as pdfium
from docx import Document
from dotenv import load_dotenv
from google.auth.transport.requests import Request
from google.oauth2.service_account import Credentials as ServiceAccountCredentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from openai import OpenAI
from watchdog.events import FileSystemEventHandler
from watchdog.observers import Observer


IGNORED_PREFIXES = (".", "~$")
IGNORED_SUFFIXES = (".crdownload", ".download", ".part", ".tmp")
SUPPORTED_SUFFIXES = {".pdf", ".docx"}
CONVERTIBLE_SUFFIXES = {".doc"}
PROCESSED_LOG_NAME = "processed_files.json"
INTAKE_LOG_NAME = "intake_log.json"
FAILED_LOG_NAME = "failed/error_log.jsonl"
IGNORED_LOG_NAME = "ignored_non_resumes/ignored_log.jsonl"
CONVERTED_DOC_ORIGINALS_DIR_NAME = "converted_doc_originals"
NO_EXTRACTED_TEXT = "No text could be extracted from resume"

WORKSHEET_COLUMNS = [
    "Rank",
    "Candidate Name",
    "Resume File",
    "Project Scope / Deliverables",
    "Stakeholder Management",
    "Project Lifecycle Ownership",
    "Years of PM Experience",
    "Qualification Total",
    "Team Size",
    "Budget Responsibility",
    "Cross-Functional Complexity",
    "Organizational Impact",
    "Scope / Scale",
    "Leadership Total",
    "Career Alignment",
    "Credential Gap",
    "Advancement Leverage",
    "Near-Term Use",
    "Career Impact Total",
    "Education",
    "Training & Certifications",
    "Professional Maturity & Communication",
    "Readiness Total",
    "Overall Score",
    "Decision",
    "Rationale",
]

SELECTION_BOARD_HEADERS = [
    "Rank",
    "Candidate ID",
    "Resume Link",
    "ID.me Validation Received",
    "Qualification Average",
    "Leadership Average",
    "Career Impact Average",
    "Readiness Average",
    "Overall Average Score",
    "Spread",
    "Confidence",
    "Selection Status",
    "Review Required",
]

CONFIDENCE_CONDITIONAL_FORMATS = [
    ("High", {"red": 0.0, "green": 0.5, "blue": 0.0}),
    ("Medium", {"red": 1.0, "green": 0.6, "blue": 0.0}),
    ("Low", {"red": 0.8, "green": 0.0, "blue": 0.0}),
]

SCORING_BOARD_HEADERS = [
    "Candidate ID",
    "Candidate Name",
    "Resume Link",
    "Pass 1 Qualification",
    "Pass 1 Leadership",
    "Pass 1 Career Impact",
    "Pass 1 Readiness",
    "Pass 1 Total",
    "Pass 1 Rationale",
    "Pass 2 Qualification",
    "Pass 2 Leadership",
    "Pass 2 Career Impact",
    "Pass 2 Readiness",
    "Pass 2 Total",
    "Pass 2 Rationale",
    "Pass 3 Qualification",
    "Pass 3 Leadership",
    "Pass 3 Career Impact",
    "Pass 3 Readiness",
    "Pass 3 Total",
    "Pass 3 Rationale",
    "Qualification Average",
    "Leadership Average",
    "Career Impact Average",
    "Readiness Average",
    "Overall Average Score",
    "Spread",
    "Confidence",
    "Timestamp",
    "Processing Status",
]

CANDIDATE_LOOKUP_HEADERS = [
    "Candidate ID",
    "Candidate Name",
    "Email",
    "Original Filename",
    "Stored Filename",
    "Resume Link",
    "Content Hash",
    "Intake Timestamp",
    "Timestamp",
    "Processing Status",
]

SCORE_KEYS = {
    "project_scope_deliverables": (0, 10),
    "stakeholder_management": (0, 10),
    "project_lifecycle_ownership": (0, 10),
    "years_pm_experience": (0, 10),
    "team_size": (0, 5),
    "budget_responsibility": (0, 5),
    "cross_functional_complexity": (0, 5),
    "organizational_impact": (0, 5),
    "scope_scale": (0, 5),
    "career_alignment": (0, 5),
    "credential_gap": (0, 5),
    "advancement_leverage": (0, 5),
    "near_term_use": (0, 5),
    "education": (0, 5),
    "training_certifications": (0, 5),
    "maturity_communication": (0, 5),
}

RESUME_SECTION_KEYWORDS = {
    "experience",
    "professional experience",
    "work experience",
    "employment",
    "education",
    "skills",
    "certifications",
    "summary",
    "professional summary",
    "projects",
    "military experience",
}

RESUME_SIGNAL_KEYWORDS = {
    "resume",
    "curriculum vitae",
    "linkedin",
    "professional summary",
    "work experience",
    "employment history",
    "project manager",
    "program manager",
    "project management",
    "stakeholder",
    "deliverables",
    "scope",
    "schedule",
    "budget",
    "cross-functional",
    "pmp",
    "scrum",
    "agile",
    "certification",
    "clearance",
}

NON_RESUME_KEYWORDS = {
    "invoice",
    "receipt",
    "statement",
    "purchase order",
    "syllabus",
    "lesson plan",
    "attendance",
    "attendee",
    "roster",
    "binder",
    "agenda",
    "minutes",
    "meeting notes",
    "contract",
    "terms and conditions",
    "privacy policy",
    "handbook",
    "flyer",
    "press release",
    "job description",
    "job posting",
    "application form",
    "cover letter",
    "transcript",
}

CONTACT_PATTERNS = [
    re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE),
    re.compile(r"\b(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b"),
    re.compile(r"\blinkedin\.com/in/[A-Za-z0-9_-]+", re.IGNORECASE),
]

EMAIL_PATTERN = CONTACT_PATTERNS[0]


@dataclass(frozen=True)
class Settings:
    openai_api_key: str
    google_credentials_json: Path
    google_token_pickle: Path
    google_sheet_id: str
    watch_folder: Path
    worksheet_name: str
    drive_upload_folder_id: str | None = None
    google_auth_mode: str = "oauth"
    openai_model: str = "gpt-4.1-mini"
    stable_checks: int = 3
    stable_interval_seconds: float = 1.0
    max_resume_chars: int = 45000
    resume_detection_mode: str = "hybrid"
    resume_confidence_threshold: int = 5
    ambiguous_confidence_min: int = 2
    allow_reprocess: bool = False
    force_new_candidate_id: bool = False
    batch_idle_minutes: float = 10.0


@dataclass(frozen=True)
class ResumeDetectionResult:
    is_resume: bool
    confidence: int
    method: str
    reasons: list[str]


def setup_logging() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(levelname)s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )


def load_settings() -> Settings:
    load_dotenv()
    missing = [
        name
        for name in [
            "OPENAI_API_KEY",
            "GOOGLE_SHEET_ID",
            "WATCH_FOLDER",
        ]
        if not os.getenv(name)
    ]
    if missing:
        raise RuntimeError(f"Missing required environment variables: {', '.join(missing)}")

    script_dir = Path(__file__).resolve().parent
    google_auth_mode = os.getenv("GOOGLE_AUTH_MODE", "oauth").lower()
    credentials_default = script_dir / "credentials.json"
    token_default = script_dir / "token.pickle"

    return Settings(
        openai_api_key=os.environ["OPENAI_API_KEY"],
        google_credentials_json=Path(
            os.getenv("GOOGLE_CREDENTIALS_JSON", str(credentials_default))
        ).expanduser().resolve(),
        google_token_pickle=Path(
            os.getenv("GOOGLE_TOKEN_PICKLE", str(token_default))
        ).expanduser().resolve(),
        google_sheet_id=os.environ["GOOGLE_SHEET_ID"],
        watch_folder=Path(os.environ["WATCH_FOLDER"]).expanduser().resolve(),
        worksheet_name=os.getenv("WORKSHEET_NAME", "Candidate Scoring Board"),
        drive_upload_folder_id=os.getenv("DRIVE_UPLOAD_FOLDER_ID") or None,
        google_auth_mode=google_auth_mode,
        openai_model=os.getenv("OPENAI_MODEL", "gpt-4.1-mini"),
        stable_checks=int(os.getenv("STABLE_CHECKS", "3")),
        stable_interval_seconds=float(os.getenv("STABLE_INTERVAL_SECONDS", "1")),
        max_resume_chars=int(os.getenv("MAX_RESUME_CHARS", "45000")),
        resume_detection_mode=os.getenv("RESUME_DETECTION_MODE", "hybrid").lower(),
        resume_confidence_threshold=int(os.getenv("RESUME_CONFIDENCE_THRESHOLD", "5")),
        ambiguous_confidence_min=int(os.getenv("AMBIGUOUS_CONFIDENCE_MIN", "2")),
        allow_reprocess=os.getenv("ALLOW_REPROCESS", "false").lower() == "true",
        force_new_candidate_id=os.getenv("FORCE_NEW_CANDIDATE_ID", "false").lower() == "true",
        batch_idle_minutes=float(os.getenv("BATCH_IDLE_MINUTES", "10")),
    )


def intake_dir(settings: Settings) -> Path:
    return settings.watch_folder / "intake"


def staging_dir(settings: Settings) -> Path:
    return settings.watch_folder / "staging"


def processed_dir(settings: Settings) -> Path:
    return settings.watch_folder / "processed"


def duplicates_dir(settings: Settings) -> Path:
    return settings.watch_folder / "duplicates"


def failed_dir(settings: Settings) -> Path:
    return settings.watch_folder / "failed"


def ignored_dir(settings: Settings) -> Path:
    return settings.watch_folder / "ignored_non_resumes"


def converted_doc_originals_dir(settings: Settings) -> Path:
    return settings.watch_folder / CONVERTED_DOC_ORIGINALS_DIR_NAME


def manifest_dir(settings: Settings) -> Path:
    return settings.watch_folder / "batch_manifests"


def intake_log_path(settings: Settings) -> Path:
    return settings.watch_folder / INTAKE_LOG_NAME


def ensure_batch_directories(settings: Settings) -> None:
    settings.watch_folder.mkdir(parents=True, exist_ok=True)
    for directory in [
        intake_dir(settings),
        staging_dir(settings),
        processed_dir(settings),
        duplicates_dir(settings),
        failed_dir(settings),
        ignored_dir(settings),
        converted_doc_originals_dir(settings),
        manifest_dir(settings),
    ]:
        directory.mkdir(exist_ok=True)


def is_supported_resume_file(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_SUFFIXES


def is_convertible_doc_file(path: Path) -> bool:
    return path.suffix.lower() in CONVERTIBLE_SUFFIXES


def should_ignore(path: Path) -> bool:
    name = path.name
    if name.startswith(IGNORED_PREFIXES):
        return True
    if name.endswith(IGNORED_SUFFIXES):
        return True
    if not is_supported_resume_file(path) and not is_convertible_doc_file(path):
        return True
    return False


def wait_until_stable(path: Path, checks: int, interval: float) -> None:
    last_size = -1
    stable_count = 0
    while stable_count < checks:
        if not path.exists():
            raise FileNotFoundError(f"File disappeared before processing: {path}")
        current_size = path.stat().st_size
        if current_size > 0 and current_size == last_size:
            stable_count += 1
        else:
            stable_count = 0
            last_size = current_size
        time.sleep(interval)


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize_resume_text_for_hash(text: str) -> str:
    normalized = text.lower()
    normalized = normalized.replace("\f", " ")
    normalized = re.sub(r"---\s*page\s+\d+\s*---", " ", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"\bpage\s+\d+\s+of\s+\d+\b", " ", normalized)
    normalized = re.sub(r"\bcreated|modified|producer|creator|metadata\b[:\s].*", " ", normalized)
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def content_hash_from_text(text: str) -> str:
    return hashlib.sha256(normalize_resume_text_for_hash(text).encode("utf-8")).hexdigest()


def content_hash_for_file(path: Path, extracted_text: str | None = None) -> str:
    if extracted_text:
        return content_hash_from_text(extracted_text)
    return file_hash(path)


def load_processed_log(log_path: Path) -> dict[str, Any]:
    if not log_path.exists():
        return {"files": {}}
    with log_path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def save_processed_log(log_path: Path, data: dict[str, Any]) -> None:
    tmp_path = log_path.with_suffix(".tmp")
    with tmp_path.open("w", encoding="utf-8") as handle:
        json.dump(data, handle, indent=2, sort_keys=True)
    tmp_path.replace(log_path)


def load_json_file(path: Path, default: dict[str, Any]) -> dict[str, Any]:
    if not path.exists():
        return default
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def save_json_file(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(exist_ok=True)
    tmp_path = path.with_suffix(".tmp")
    with tmp_path.open("w", encoding="utf-8") as handle:
        json.dump(data, handle, indent=2, sort_keys=True)
    tmp_path.replace(path)


def file_timestamp(path: Path) -> str:
    return datetime.fromtimestamp(path.stat().st_mtime, timezone.utc).isoformat()


def record_intake_file(settings: Settings, path: Path) -> None:
    if path.parent != intake_dir(settings) or should_ignore(path) or not path.exists():
        return
    log_path = intake_log_path(settings)
    intake_log = load_json_file(log_path, {"files": {}})
    key = path.name
    intake_log.setdefault("files", {}).setdefault(
        key,
        {
            "original_filename": path.name,
            "intake_timestamp": file_timestamp(path),
            "first_seen_at": datetime.now(timezone.utc).isoformat(),
        },
    )
    save_json_file(log_path, intake_log)
    logging.info("Staged intake metadata for %s; no scoring performed", path.name)


def intake_timestamp_for(settings: Settings, path: Path) -> str:
    intake_log = load_json_file(intake_log_path(settings), {"files": {}})
    entry = intake_log.get("files", {}).get(path.name)
    if entry and entry.get("intake_timestamp"):
        return str(entry["intake_timestamp"])
    return file_timestamp(path)


def remove_intake_log_entry(settings: Settings, filename: str) -> None:
    log_path = intake_log_path(settings)
    intake_log = load_json_file(log_path, {"files": {}})
    files = intake_log.setdefault("files", {})
    if filename in files:
        del files[filename]
        save_json_file(log_path, intake_log)


def append_failure_log(watch_folder: Path, path: Path, reason: str) -> None:
    failed_dir = watch_folder / "failed"
    failed_dir.mkdir(exist_ok=True)
    log_path = watch_folder / FAILED_LOG_NAME
    entry = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "file": str(path),
        "reason": reason,
    }
    with log_path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(entry) + "\n")


def append_ignored_log(
    watch_folder: Path,
    path: Path,
    detection: ResumeDetectionResult,
) -> None:
    ignored_dir = watch_folder / "ignored_non_resumes"
    ignored_dir.mkdir(exist_ok=True)
    log_path = watch_folder / IGNORED_LOG_NAME
    entry = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "file": str(path),
        "confidence": detection.confidence,
        "method": detection.method,
        "reasons": detection.reasons,
    }
    with log_path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(entry) + "\n")


def unique_destination(directory: Path, filename: str) -> Path:
    directory.mkdir(exist_ok=True)
    candidate = directory / filename
    if not candidate.exists():
        return candidate

    stem = candidate.stem
    suffix = candidate.suffix
    for index in range(1, 10_000):
        next_candidate = directory / f"{stem}_{index}{suffix}"
        if not next_candidate.exists():
            return next_candidate
    raise RuntimeError(f"Could not create unique destination for {filename}")


def move_file(path: Path, destination_dir: Path) -> Path:
    destination = unique_destination(destination_dir, path.name)
    shutil.move(str(path), str(destination))
    return destination


def convert_doc_to_docx(settings: Settings, path: Path) -> Path:
    """Convert legacy .doc files without requiring a human to open/read them."""
    if not is_convertible_doc_file(path):
        return path

    wait_until_stable(path, settings.stable_checks, settings.stable_interval_seconds)
    destination = unique_destination(intake_dir(settings), f"{path.stem}.docx")
    tmp_destination = destination.with_suffix(".docx.tmp")
    if tmp_destination.exists():
        tmp_destination.unlink()

    try:
        subprocess.run(
            [
                "textutil",
                "-convert",
                "docx",
                "-output",
                str(tmp_destination),
                str(path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        tmp_destination.replace(destination)
    except Exception:
        if tmp_destination.exists():
            tmp_destination.unlink()
        raise

    original_destination = unique_destination(converted_doc_originals_dir(settings), path.name)
    shutil.move(str(path), str(original_destination))
    logging.info(
        "Converted legacy Word file to intake DOCX: %s -> %s; archived original at %s",
        path.name,
        destination.name,
        original_destination,
    )
    record_intake_file(settings, destination)
    return destination


def extract_pdf_text(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"\n--- Page {page_number} ---\n{text}")
    return "\n".join(parts).strip()


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_resume_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = extract_pdf_text(path)
    elif suffix == ".docx":
        text = extract_docx_text(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    if not text:
        raise ValueError(NO_EXTRACTED_TEXT)
    return text


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower()).strip()


def count_keyword_hits(text: str, keywords: set[str]) -> list[str]:
    normalized = normalize_text(text)
    return sorted(keyword for keyword in keywords if keyword in normalized)


def contact_hits(text: str) -> list[str]:
    hits: list[str] = []
    for pattern in CONTACT_PATTERNS:
        if pattern.search(text):
            hits.append(pattern.pattern)
    return hits


def extract_email(text: str) -> str:
    match = EMAIL_PATTERN.search(text)
    return match.group(0).strip().lower() if match else ""


def valid_email(value: str) -> str:
    return extract_email(value)


def filename_resume_hints(path: Path) -> list[str]:
    normalized = normalize_text(path.stem.replace("_", " ").replace("-", " "))
    hints = []
    for keyword in ["resume", "cv", "curriculum vitae", "professional profile"]:
        if keyword in normalized:
            hints.append(keyword)
    return hints


def heuristic_resume_detection(
    path: Path,
    text: str,
    threshold: int,
) -> ResumeDetectionResult:
    section_hits = count_keyword_hits(text, RESUME_SECTION_KEYWORDS)
    signal_hits = count_keyword_hits(text, RESUME_SIGNAL_KEYWORDS)
    negative_hits = count_keyword_hits(text, NON_RESUME_KEYWORDS)
    contacts = contact_hits(text)
    filename_hits = filename_resume_hints(path)

    score = 0
    reasons: list[str] = []

    if filename_hits:
        score += 3
        reasons.append(f"filename hints: {', '.join(filename_hits)}")
    if contacts:
        score += min(len(contacts), 3)
        reasons.append(f"contact signals: {len(contacts)}")
    if section_hits:
        score += min(len(section_hits), 5)
        reasons.append(f"resume sections: {', '.join(section_hits[:8])}")
    if signal_hits:
        score += min(len(signal_hits), 5)
        reasons.append(f"resume keywords: {', '.join(signal_hits[:8])}")
    if len(text) >= 1200:
        score += 1
        reasons.append("resume-length document")
    if negative_hits:
        penalty = min(len(negative_hits), 6)
        score -= penalty
        reasons.append(f"non-resume signals: {', '.join(negative_hits[:8])}")

    confidence = max(score, 0)
    return ResumeDetectionResult(
        is_resume=confidence >= threshold,
        confidence=confidence,
        method="heuristic",
        reasons=reasons or ["no strong resume signals found"],
    )


def build_resume_detection_prompt(filename: str, resume_text: str) -> str:
    return f"""
Decide whether this document is a resume/CV for an individual candidate.

Return strict JSON only with this exact shape:
{{
  "is_resume": true,
  "confidence": 0,
  "reasons": []
}}

Rules:
- is_resume must be true only for resumes/CVs/professional profiles.
- Cover letters, applications, invoices, attendance lists, syllabi, meeting notes, binders, job postings, and transcripts are not resumes.
- confidence must be an integer from 0 to 10.
- reasons must be short strings based only on evidence in the document.

Filename: {filename}

Document text:
{resume_text}
""".strip()


def llm_resume_detection(
    client: OpenAI,
    model: str,
    path: Path,
    text: str,
    max_chars: int,
) -> ResumeDetectionResult:
    prompt = build_resume_detection_prompt(path.name, text[:max_chars])
    response = client.chat.completions.create(
        model=model,
        temperature=0,
        messages=[
            {"role": "system", "content": "You classify documents. Return strict JSON only."},
            {"role": "user", "content": prompt},
        ],
    )
    content = response.choices[0].message.content or ""
    parsed = extract_json_object(content)
    is_resume = parsed.get("is_resume")
    confidence = parsed.get("confidence")
    reasons = parsed.get("reasons")
    if not isinstance(is_resume, bool):
        raise ValueError("Resume detection response missing boolean is_resume")
    if not isinstance(confidence, int) or confidence < 0 or confidence > 10:
        raise ValueError("Resume detection confidence must be an integer from 0 to 10")
    if not isinstance(reasons, list) or not all(isinstance(reason, str) for reason in reasons):
        raise ValueError("Resume detection reasons must be a list of strings")
    return ResumeDetectionResult(
        is_resume=is_resume and confidence >= 6,
        confidence=confidence,
        method="llm",
        reasons=reasons,
    )


def detect_resume(
    settings: Settings,
    client: OpenAI,
    path: Path,
    text: str,
) -> ResumeDetectionResult:
    heuristic = heuristic_resume_detection(path, text, settings.resume_confidence_threshold)
    mode = settings.resume_detection_mode
    if mode == "heuristic":
        return heuristic
    if mode not in {"hybrid", "llm"}:
        raise ValueError("RESUME_DETECTION_MODE must be heuristic, hybrid, or llm")
    if mode == "hybrid" and heuristic.confidence >= settings.resume_confidence_threshold:
        return heuristic
    if mode == "hybrid" and heuristic.confidence < settings.ambiguous_confidence_min:
        return heuristic

    logging.info(
        "Resume detection is ambiguous for %s; asking LLM classifier. Heuristic confidence=%s",
        path.name,
        heuristic.confidence,
    )
    llm_result = llm_resume_detection(
        client,
        settings.openai_model,
        path,
        text,
        min(settings.max_resume_chars, 12000),
    )
    return ResumeDetectionResult(
        is_resume=llm_result.is_resume,
        confidence=llm_result.confidence,
        method=f"{heuristic.method}+{llm_result.method}",
        reasons=[*heuristic.reasons, *llm_result.reasons],
    )


def build_scoring_prompt(resume_text: str) -> str:
    return f"""
You are scoring one PMP scholarship/course applicant using the rubric below.

Return strict JSON only. Do not include markdown, commentary, or extra keys.
Use integer scores only. Stay within each score range.

Qualification Evidence: 40 points
- project_scope_deliverables: 0-10
- stakeholder_management: 0-10
- project_lifecycle_ownership: 0-10
- years_pm_experience: 0-10

Leadership & Complexity: 25 points
- team_size: 0-5
- budget_responsibility: 0-5
- cross_functional_complexity: 0-5
- organizational_impact: 0-5
- scope_scale: 0-5

Career Impact: 20 points
- career_alignment: 0-5
- credential_gap: 0-5
- advancement_leverage: 0-5
- near_term_use: 0-5

Readiness: 15 points
- education: 0-5
- training_certifications: 0-5
- maturity_communication: 0-5

Required JSON shape:
{{
  "candidate_name": "",
  "candidate_email": "",
  "scores": {{
    "project_scope_deliverables": 0,
    "stakeholder_management": 0,
    "project_lifecycle_ownership": 0,
    "years_pm_experience": 0,
    "team_size": 0,
    "budget_responsibility": 0,
    "cross_functional_complexity": 0,
    "organizational_impact": 0,
    "scope_scale": 0,
    "career_alignment": 0,
    "credential_gap": 0,
    "advancement_leverage": 0,
    "near_term_use": 0,
    "education": 0,
    "training_certifications": 0,
    "maturity_communication": 0
  }},
  "rationale": ""
}}

Resume text:
{resume_text}
""".strip()


def extract_json_object(raw_text: str) -> dict[str, Any]:
    cleaned = raw_text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    return json.loads(cleaned)


def score_resume(
    client: OpenAI,
    model: str,
    resume_text: str,
    max_chars: int,
    review_pass: int | None = None,
) -> dict[str, Any]:
    prompt = build_scoring_prompt(resume_text[:max_chars])
    if review_pass is not None:
        prompt = (
            f"This is independent scoring pass {review_pass} of 3. "
            "Score from the evidence only; do not try to match any other pass.\n\n"
            f"{prompt}"
        )
    response = client.chat.completions.create(
        model=model,
        temperature=0.15 if review_pass is not None else 0,
        messages=[
            {"role": "system", "content": "You are a precise resume scoring assistant. Return strict JSON only."},
            {"role": "user", "content": prompt},
        ],
    )
    content = response.choices[0].message.content or ""
    parsed = extract_json_object(content)
    validate_score_payload(parsed)
    return parsed


def render_pdf_pages_as_data_urls(path: Path, max_pages: int = 5, scale: float = 2.0) -> list[str]:
    data_urls: list[str] = []
    pdf = pdfium.PdfDocument(str(path))
    try:
        page_count = min(len(pdf), max_pages)
        for index in range(page_count):
            page = pdf[index]
            pil_image = page.render(scale=scale).to_pil()
            if pil_image.mode != "RGB":
                pil_image = pil_image.convert("RGB")
            # Keep requests reasonably sized while preserving readable resume text.
            pil_image.thumbnail((1800, 2400))
            buffer = io.BytesIO()
            pil_image.save(buffer, format="JPEG", quality=85)
            encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
            data_urls.append(f"data:image/jpeg;base64,{encoded}")
    finally:
        pdf.close()
    if not data_urls:
        raise ValueError(f"Could not render PDF pages for {path.name}")
    return data_urls


def build_vision_scoring_prompt(filename: str) -> str:
    return f"""
The attached images are pages from a PDF that did not contain extractable text.
If this is a resume/CV/professional profile, score it for PMP selection.

Return strict JSON only. Do not include markdown, commentary, or extra keys.
Use integer scores only. Stay within each score range.
Infer the candidate name from the resume. If the document is not a resume, return
the best available candidate_name as an empty string and explain that it is not a
resume in rationale.

Filename: {filename}

Qualification Evidence: 40 points
- project_scope_deliverables: 0-10
- stakeholder_management: 0-10
- project_lifecycle_ownership: 0-10
- years_pm_experience: 0-10

Leadership & Complexity: 25 points
- team_size: 0-5
- budget_responsibility: 0-5
- cross_functional_complexity: 0-5
- organizational_impact: 0-5
- scope_scale: 0-5

Career Impact: 20 points
- career_alignment: 0-5
- credential_gap: 0-5
- advancement_leverage: 0-5
- near_term_use: 0-5

Readiness: 15 points
- education: 0-5
- training_certifications: 0-5
- maturity_communication: 0-5

Required JSON shape:
{{
  "candidate_name": "",
  "candidate_email": "",
  "scores": {{
    "project_scope_deliverables": 0,
    "stakeholder_management": 0,
    "project_lifecycle_ownership": 0,
    "years_pm_experience": 0,
    "team_size": 0,
    "budget_responsibility": 0,
    "cross_functional_complexity": 0,
    "organizational_impact": 0,
    "scope_scale": 0,
    "career_alignment": 0,
    "credential_gap": 0,
    "advancement_leverage": 0,
    "near_term_use": 0,
    "education": 0,
    "training_certifications": 0,
    "maturity_communication": 0
  }},
  "rationale": ""
}}
""".strip()


def score_pdf_resume_from_images(
    client: OpenAI,
    model: str,
    path: Path,
    review_pass: int | None = None,
) -> dict[str, Any]:
    data_urls = render_pdf_pages_as_data_urls(path)
    prompt = build_vision_scoring_prompt(path.name)
    if review_pass is not None:
        prompt = (
            f"This is independent scoring pass {review_pass} of 3. "
            "Score from the evidence only; do not try to match any other pass.\n\n"
            f"{prompt}"
        )
    content: list[dict[str, Any]] = [
        {"type": "text", "text": prompt}
    ]
    content.extend(
        {"type": "image_url", "image_url": {"url": data_url}}
        for data_url in data_urls
    )
    response = client.chat.completions.create(
        model=model,
        temperature=0.15 if review_pass is not None else 0,
        messages=[
            {
                "role": "system",
                "content": "You are a precise resume scoring assistant. Return strict JSON only.",
            },
            {"role": "user", "content": content},
        ],
    )
    content_text = response.choices[0].message.content or ""
    parsed = extract_json_object(content_text)
    validate_score_payload(parsed)
    return parsed


def validate_score_payload(payload: dict[str, Any]) -> None:
    if not isinstance(payload.get("candidate_name"), str) or not payload["candidate_name"].strip():
        raise ValueError("LLM response missing candidate_name")
    if not isinstance(payload.get("rationale"), str):
        raise ValueError("LLM response missing rationale")
    scores = payload.get("scores")
    if not isinstance(scores, dict):
        raise ValueError("LLM response missing scores object")

    missing = sorted(set(SCORE_KEYS) - set(scores))
    extra = sorted(set(scores) - set(SCORE_KEYS))
    if missing or extra:
        raise ValueError(f"Score keys mismatch. Missing={missing}; Extra={extra}")

    for key, (minimum, maximum) in SCORE_KEYS.items():
        value = scores[key]
        if not isinstance(value, int):
            raise ValueError(f"{key} must be an integer")
        if value < minimum or value > maximum:
            raise ValueError(f"{key}={value} outside allowed range {minimum}-{maximum}")


def google_scopes(settings: Settings) -> list[str]:
    scopes = ["https://www.googleapis.com/auth/spreadsheets"]
    if settings.drive_upload_folder_id:
        scopes.append("https://www.googleapis.com/auth/drive")
    return scopes


def get_google_credentials(settings: Settings):
    scopes = google_scopes(settings)
    if settings.google_auth_mode == "service_account":
        return ServiceAccountCredentials.from_service_account_file(
            settings.google_credentials_json,
            scopes=scopes,
        )
    if settings.google_auth_mode == "oauth":
        return get_oauth_credentials(settings, scopes)
    raise ValueError("GOOGLE_AUTH_MODE must be oauth or service_account")


def get_worksheet(settings: Settings, creds) -> gspread.Worksheet:
    client = gspread.authorize(creds)
    spreadsheet = client.open_by_key(settings.google_sheet_id)
    worksheet = spreadsheet.worksheet(settings.worksheet_name)
    return worksheet


def get_oauth_credentials(settings: Settings, scopes: list[str]):
    import pickle

    creds = None
    if settings.google_token_pickle.exists():
        with settings.google_token_pickle.open("rb") as token:
            creds = pickle.load(token)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not settings.google_credentials_json.exists():
                raise FileNotFoundError(
                    f"OAuth credentials file not found: {settings.google_credentials_json}"
                )
            flow = InstalledAppFlow.from_client_secrets_file(
                str(settings.google_credentials_json),
                scopes,
            )
            creds = flow.run_local_server(port=0)
        with settings.google_token_pickle.open("wb") as token:
            pickle.dump(creds, token)

    return creds


def ensure_headers(worksheet: gspread.Worksheet) -> None:
    existing = worksheet.row_values(1)
    if existing[: len(WORKSHEET_COLUMNS)] == WORKSHEET_COLUMNS:
        return
    logging.info("Updating worksheet headers to expected PMP scoring layout")
    worksheet.update("A1:Z1", [WORKSHEET_COLUMNS])


def first_empty_candidate_row(worksheet: gspread.Worksheet) -> int:
    names = worksheet.col_values(2)
    return max(len(names) + 1, 2)


def sheet_string(value: str) -> str:
    return value.replace('"', '""')


def resume_file_cell_value(filename: str, drive_link: str | None) -> str:
    if not drive_link:
        return filename
    return f'=HYPERLINK("{sheet_string(drive_link)}","{sheet_string(filename)}")'


def ensure_row_headers(worksheet: gspread.Worksheet, headers: list[str]) -> None:
    existing = worksheet.row_values(1)
    if existing[: len(headers)] != headers:
        worksheet.update(f"A1:{column_letter(len(headers))}1", [headers])


def ensure_selection_confidence_formatting(selection_worksheet: gspread.Worksheet) -> None:
    sheet_id = selection_worksheet.id
    confidence_range = {
        "sheetId": sheet_id,
        "startRowIndex": 1,
        "endRowIndex": 500,
        "startColumnIndex": 10,
        "endColumnIndex": 11,
    }
    metadata = selection_worksheet.spreadsheet.fetch_sheet_metadata(
        params={"includeGridData": False}
    )
    existing_rule_indexes = []
    for sheet in metadata.get("sheets", []):
        properties = sheet.get("properties", {})
        if properties.get("sheetId") != sheet_id:
            continue
        rules = sheet.get("conditionalFormats", [])
        for index, rule in enumerate(rules):
            ranges = rule.get("ranges", [])
            condition = rule.get("booleanRule", {}).get("condition", {})
            values = {
                value.get("userEnteredValue")
                for value in condition.get("values", [])
            }
            if (
                len(ranges) == 1
                and ranges[0] == confidence_range
                and condition.get("type") == "TEXT_EQ"
                and values <= {value for value, _color in CONFIDENCE_CONDITIONAL_FORMATS}
            ):
                existing_rule_indexes.append(index)

    requests = [
        {
            "deleteConditionalFormatRule": {
                "sheetId": sheet_id,
                "index": index,
            }
        }
        for index in sorted(existing_rule_indexes, reverse=True)
    ]
    for value, color in reversed(CONFIDENCE_CONDITIONAL_FORMATS):
        requests.append(
            {
                "addConditionalFormatRule": {
                    "rule": {
                        "ranges": [confidence_range],
                        "booleanRule": {
                            "condition": {
                                "type": "TEXT_EQ",
                                "values": [{"userEnteredValue": value}],
                            },
                            "format": {
                                "textFormat": {
                                    "foregroundColor": color,
                                    "bold": True,
                                }
                            },
                        },
                    },
                    "index": 0,
                }
            }
        )

    if requests:
        selection_worksheet.spreadsheet.batch_update({"requests": requests})


def column_letter(index_1_based: int) -> str:
    letters = ""
    n = index_1_based
    while n:
        n, rem = divmod(n - 1, 26)
        letters = chr(65 + rem) + letters
    return letters


def next_candidate_id(lookup_worksheet: gspread.Worksheet) -> str:
    existing_ids = lookup_worksheet.col_values(1)[1:]
    max_number = 0
    for candidate_id in existing_ids:
        match = re.fullmatch(r"C-?(\d{4})", candidate_id.strip())
        if match:
            max_number = max(max_number, int(match.group(1)))
    return f"C{max_number + 1:04d}"


def candidate_id_from_number(number: int) -> str:
    return f"C{number:04d}"


def candidate_id_number(candidate_id: str) -> int:
    match = re.fullmatch(r"C-?(\d{4})", candidate_id.strip())
    if not match:
        raise ValueError(f"Unexpected candidate ID format: {candidate_id}")
    return int(match.group(1))


def next_batch_id(settings: Settings) -> str:
    prefix = datetime.now().strftime("PMP-%Y-%m-%d")
    existing_numbers = []
    for path in manifest_dir(settings).glob(f"{prefix}-*.json"):
        match = re.fullmatch(rf"{re.escape(prefix)}-(\d{{3}})", path.stem)
        if match:
            existing_numbers.append(int(match.group(1)))
    return f"{prefix}-{(max(existing_numbers) if existing_numbers else 0) + 1:03d}"


def find_lookup_by_hash(lookup_worksheet: gspread.Worksheet, content_hash: str) -> dict[str, Any] | None:
    rows = lookup_worksheet.get_all_values()
    if not rows:
        return None
    headers = rows[0]

    def value(row: list[str], header: str, fallback_index: int | None = None) -> str:
        try:
            index = headers.index(header)
        except ValueError:
            index = fallback_index
        if index is None or len(row) <= index:
            return ""
        return row[index]

    for row_number, row in enumerate(rows[1:], start=2):
        row_content_hash = value(row, "Content Hash", 5).strip()
        if row_content_hash == content_hash:
            return {
                "row": row_number,
                "candidate_id": value(row, "Candidate ID", 0),
                "candidate_name": value(row, "Candidate Name", 1),
                "email": value(row, "Email"),
                "original_filename": value(row, "Original Filename", 2),
                "stored_filename": value(row, "Stored Filename", 3),
                "resume_link": value(row, "Resume Link", 4),
                "content_hash": row_content_hash,
                "intake_timestamp": value(row, "Intake Timestamp", 6),
                "status": value(row, "Processing Status", 8) or value(row, "Timestamp", 7),
            }
    return None


def find_scoring_row_by_candidate_id(scoring_worksheet: gspread.Worksheet, candidate_id: str) -> int | None:
    ids = scoring_worksheet.col_values(1)
    for row_number, value in enumerate(ids, start=1):
        if value.strip() == candidate_id:
            return row_number
    return None


def category_totals(payload: dict[str, Any]) -> dict[str, Any]:
    scores = payload["scores"]
    qualification = sum(
        scores[key]
        for key in [
            "project_scope_deliverables",
            "stakeholder_management",
            "project_lifecycle_ownership",
            "years_pm_experience",
        ]
    )
    leadership = sum(
        scores[key]
        for key in [
            "team_size",
            "budget_responsibility",
            "cross_functional_complexity",
            "organizational_impact",
            "scope_scale",
        ]
    )
    career = sum(
        scores[key]
        for key in [
            "career_alignment",
            "credential_gap",
            "advancement_leverage",
            "near_term_use",
        ]
    )
    readiness = sum(
        scores[key]
        for key in [
            "education",
            "training_certifications",
            "maturity_communication",
        ]
    )
    total = qualification + leadership + career + readiness
    return {
        "candidate_name": payload["candidate_name"].strip(),
        "candidate_email": valid_email(payload.get("candidate_email", "")),
        "qualification": qualification,
        "leadership": leadership,
        "career": career,
        "readiness": readiness,
        "total": total,
        "rationale": payload["rationale"].strip(),
    }


def average(values: list[float]) -> float:
    return round(sum(values) / len(values), 1)


def confidence_from_spread(spread: float) -> str:
    if spread <= 2:
        return "High"
    if spread <= 5:
        return "Medium"
    return "Low"


def score_resume_three_passes(
    client: OpenAI,
    model: str,
    resume_text: str,
    max_chars: int,
) -> list[dict[str, Any]]:
    passes = []
    for pass_number in range(1, 4):
        logging.info("Starting AI scoring pass %s/3", pass_number)
        passes.append(category_totals(score_resume(client, model, resume_text, max_chars, pass_number)))
    return passes


def score_pdf_resume_three_passes_from_images(
    client: OpenAI,
    model: str,
    path: Path,
) -> list[dict[str, Any]]:
    passes = []
    for pass_number in range(1, 4):
        logging.info("Starting rendered-PDF AI scoring pass %s/3", pass_number)
        passes.append(category_totals(score_pdf_resume_from_images(client, model, path, pass_number)))
    return passes


def summarize_passes(passes: list[dict[str, Any]]) -> dict[str, Any]:
    totals = [item["total"] for item in passes]
    spread = max(totals) - min(totals)
    return {
        "qualification_average": average([item["qualification"] for item in passes]),
        "leadership_average": average([item["leadership"] for item in passes]),
        "career_average": average([item["career"] for item in passes]),
        "readiness_average": average([item["readiness"] for item in passes]),
        "overall_average": average(totals),
        "spread": spread,
        "confidence": confidence_from_spread(spread),
    }


def append_v3_candidate(
    scoring_worksheet: gspread.Worksheet,
    lookup_worksheet: gspread.Worksheet,
    candidate_id: str,
    candidate_name: str,
    email: str,
    original_filename: str,
    stored_filename: str,
    resume_link: str,
    content_hash: str,
    intake_timestamp: str,
    passes: list[dict[str, Any]],
    summary: dict[str, Any],
    scoring_row: int | None = None,
    lookup_row: int | None = None,
) -> None:
    timestamp = datetime.now(timezone.utc).isoformat()
    link_formula = resume_file_cell_value(stored_filename, resume_link)
    lookup_values = [
        candidate_id,
        candidate_name,
        email,
        original_filename,
        stored_filename,
        link_formula,
        content_hash,
        intake_timestamp,
        timestamp,
        "Processed",
    ]
    scoring_values = [
        candidate_id,
        candidate_name,
        link_formula,
        passes[0]["qualification"],
        passes[0]["leadership"],
        passes[0]["career"],
        passes[0]["readiness"],
        passes[0]["total"],
        passes[0]["rationale"],
        passes[1]["qualification"],
        passes[1]["leadership"],
        passes[1]["career"],
        passes[1]["readiness"],
        passes[1]["total"],
        passes[1]["rationale"],
        passes[2]["qualification"],
        passes[2]["leadership"],
        passes[2]["career"],
        passes[2]["readiness"],
        passes[2]["total"],
        passes[2]["rationale"],
        summary["qualification_average"],
        summary["leadership_average"],
        summary["career_average"],
        summary["readiness_average"],
        summary["overall_average"],
        summary["spread"],
        summary["confidence"],
        timestamp,
        "Processed",
    ]
    if lookup_row:
        lookup_worksheet.update(f"A{lookup_row}:J{lookup_row}", [lookup_values], value_input_option="USER_ENTERED")
    else:
        lookup_worksheet.append_row(lookup_values, value_input_option="USER_ENTERED")

    if scoring_row:
        scoring_worksheet.update(f"A{scoring_row}:AD{scoring_row}", [scoring_values], value_input_option="USER_ENTERED")
    else:
        scoring_worksheet.append_row(scoring_values, value_input_option="USER_ENTERED")


def selection_status(rank: int, available_seats: int) -> str:
    return "Selected" if rank <= available_seats else "Waitlist"


def selection_status_formula(row_number: int) -> str:
    return f'=IF($A{row_number}="","",IF($A{row_number}<=$Q$4,"Selected","Waitlist"))'


def rebuild_selection_board(selection_worksheet: gspread.Worksheet, scoring_worksheet: gspread.Worksheet) -> list[dict[str, Any]]:
    ensure_row_headers(selection_worksheet, SELECTION_BOARD_HEADERS)
    ensure_selection_confidence_formatting(selection_worksheet)
    existing_idme_values = {}
    for row in selection_worksheet.get("B2:D500", value_render_option="FORMATTED_VALUE"):
        if len(row) >= 3 and row[0].strip() and row[2].strip() in {"Yes", "No"}:
            existing_idme_values[row[0].strip()] = row[2].strip()
    scoring_rows = scoring_worksheet.get("A2:AD500", value_render_option="FORMULA")
    candidates = []
    for row in scoring_rows:
        if len(row) < 28 or not row[0].strip():
            continue
        candidates.append(
            {
                "candidate_id": row[0],
                "resume_link": row[2] if len(row) > 2 else "",
                "qualification_average": row[21],
                "leadership_average": row[22],
                "career_average": row[23],
                "readiness_average": row[24],
                "overall_average": float(row[25]),
                "spread": float(row[26]),
                "confidence": row[27],
            }
        )
    candidates.sort(key=lambda item: item["overall_average"], reverse=True)
    settings = selection_worksheet.get("P2:Q4")
    available_seats = 18
    for setting in settings:
        if len(setting) >= 2 and setting[0] == "Available Seats":
            try:
                available_seats = int(float(setting[1]))
            except ValueError:
                available_seats = 18

    output = []
    summaries = []
    for rank, candidate in enumerate(candidates, start=1):
        row_number = rank + 1
        review_required = "Yes" if candidate["spread"] > 5 or candidate["confidence"] == "Low" else "No"
        status = selection_status(rank, available_seats)
        output.append(
            [
                rank,
                candidate["candidate_id"],
                candidate["resume_link"],
                existing_idme_values.get(candidate["candidate_id"], "No"),
                candidate["qualification_average"],
                candidate["leadership_average"],
                candidate["career_average"],
                candidate["readiness_average"],
                candidate["overall_average"],
                candidate["spread"],
                candidate["confidence"],
                selection_status_formula(row_number),
                review_required,
            ]
        )
        summaries.append({"candidate_id": candidate["candidate_id"], "rank": rank, "overall_average": candidate["overall_average"], "status": status})

    selection_worksheet.batch_clear(["A2:M500"])
    if output:
        selection_worksheet.update(f"A2:M{len(output) + 1}", output, value_input_option="USER_ENTERED")
    return summaries


def append_to_sheet(
    worksheet: gspread.Worksheet,
    filename: str,
    payload: dict[str, Any],
    drive_link: str | None = None,
) -> int:
    row_number = first_empty_candidate_row(worksheet)
    scores = payload["scores"]
    worksheet.batch_update(
        [
            {
                "range": f"B{row_number}:G{row_number}",
                "values": [[
                    payload["candidate_name"].strip(),
                    resume_file_cell_value(filename, drive_link),
                    scores["project_scope_deliverables"],
                    scores["stakeholder_management"],
                    scores["project_lifecycle_ownership"],
                    scores["years_pm_experience"],
                ]],
            },
            {
                "range": f"I{row_number}:M{row_number}",
                "values": [[
                    scores["team_size"],
                    scores["budget_responsibility"],
                    scores["cross_functional_complexity"],
                    scores["organizational_impact"],
                    scores["scope_scale"],
                ]],
            },
            {
                "range": f"O{row_number}:R{row_number}",
                "values": [[
                    scores["career_alignment"],
                    scores["credential_gap"],
                    scores["advancement_leverage"],
                    scores["near_term_use"],
                ]],
            },
            {
                "range": f"T{row_number}:V{row_number}",
                "values": [[
                    scores["education"],
                    scores["training_certifications"],
                    scores["maturity_communication"],
                ]],
            },
            {
                "range": f"Z{row_number}",
                "values": [[payload["rationale"].strip()]],
            },
        ],
        value_input_option="USER_ENTERED",
    )
    logging.info("Appended %s to row %s", payload["candidate_name"], row_number)
    return row_number


def upload_resume_copy(
    drive_service,
    path: Path,
    folder_id: str,
    stored_filename: str | None = None,
    attempts: int = 3,
) -> dict[str, str]:
    metadata = {
        "name": stored_filename or path.name,
        "parents": [folder_id],
    }
    last_error: Exception | None = None
    for attempt in range(1, attempts + 1):
        try:
            media = MediaFileUpload(str(path), resumable=True)
            uploaded = drive_service.files().create(
                body=metadata,
                media_body=media,
                fields="id,name,webViewLink",
            ).execute()
            logging.info(
                "Uploaded resume copy to Drive: %s (%s)",
                uploaded.get("name"),
                uploaded.get("webViewLink") or uploaded.get("id"),
            )
            return uploaded
        except Exception as exc:  # noqa: BLE001
            last_error = exc
            if attempt == attempts:
                break
            wait_seconds = attempt * 2
            logging.warning(
                "Drive upload failed for %s on attempt %s/%s: %s. Retrying in %ss",
                path.name,
                attempt,
                attempts,
                exc,
                wait_seconds,
            )
            time.sleep(wait_seconds)
    raise RuntimeError(f"Drive upload failed after {attempts} attempts: {last_error}")


def sort_sheet_by_rank(sheets_service, spreadsheet_id: str, sheet_id: int) -> None:
    sheets_service.spreadsheets().batchUpdate(
        spreadsheetId=spreadsheet_id,
        body={
            "requests": [
                {
                    "sortRange": {
                        "range": {
                            "sheetId": sheet_id,
                            "startRowIndex": 1,
                            "endRowIndex": 100,
                            "startColumnIndex": 0,
                            "endColumnIndex": 27,
                        },
                        "sortSpecs": [
                            {"dimensionIndex": 0, "sortOrder": "ASCENDING"}
                        ],
                    }
                }
            ]
        },
    ).execute()
    logging.info("Sorted sheet by Rank column")


def get_processed_result_summary(
    worksheet: gspread.Worksheet,
    candidate_name: str,
    filename: str,
) -> dict[str, str] | None:
    rows = worksheet.get("A2:Z100", value_render_option="FORMATTED_VALUE")
    normalized_candidate = candidate_name.strip().lower()
    normalized_filename = filename.strip().lower()
    fallback_match = None

    for index, row in enumerate(rows, start=2):
        row_candidate = (row[1] if len(row) > 1 else "").strip().lower()
        row_filename = (row[2] if len(row) > 2 else "").strip().lower()
        if row_candidate != normalized_candidate:
            continue

        summary = {
            "row": str(index),
            "rank": row[0] if len(row) > 0 else "",
            "candidate": row[1] if len(row) > 1 else candidate_name,
            "resume_file": row[2] if len(row) > 2 else filename,
            "overall_score": row[23] if len(row) > 23 else "",
            "decision": row[24] if len(row) > 24 else "",
        }
        if row_filename == normalized_filename:
            return summary
        if fallback_match is None:
            fallback_match = summary

    return fallback_match


def log_processed_result_summary(summary: dict[str, str] | None) -> None:
    if not summary:
        logging.warning("Could not find the processed candidate row after sorting")
        return

    logging.info(
        "LATEST RESULT | Candidate: %s | Rank: %s | Total Score: %s | Decision: %s | Row: %s",
        summary["candidate"],
        summary["rank"],
        summary["overall_score"],
        summary["decision"],
        summary["row"],
    )


class ResumeHandler(FileSystemEventHandler):
    def __init__(self, settings: Settings):
        self.settings = settings
        self.client = OpenAI(api_key=settings.openai_api_key)
        self.google_creds = get_google_credentials(settings)
        self.worksheet = get_worksheet(settings, self.google_creds)
        spreadsheet = self.worksheet.spreadsheet
        self.selection_worksheet = spreadsheet.worksheet("Selection Board")
        self.lookup_worksheet = spreadsheet.worksheet("Candidate Lookup")
        ensure_row_headers(self.selection_worksheet, SELECTION_BOARD_HEADERS)
        ensure_selection_confidence_formatting(self.selection_worksheet)
        ensure_row_headers(self.worksheet, SCORING_BOARD_HEADERS)
        ensure_row_headers(self.lookup_worksheet, CANDIDATE_LOOKUP_HEADERS)
        self.drive_service = (
            build("drive", "v3", credentials=self.google_creds)
            if settings.drive_upload_folder_id
            else None
        )
        self.sheets_service = build("sheets", "v4", credentials=self.google_creds)
        self.processed_log_path = settings.watch_folder / PROCESSED_LOG_NAME
        self.processed_log = load_processed_log(self.processed_log_path)

    def try_process(
        self,
        path: Path,
        candidate_id: str | None = None,
        original_filename: str | None = None,
        intake_timestamp: str | None = None,
        batch_id: str | None = None,
    ) -> dict[str, Any]:
        try:
            return self.process(
                path,
                candidate_id=candidate_id,
                original_filename=original_filename,
                intake_timestamp=intake_timestamp,
                batch_id=batch_id,
            )
        except Exception as exc:  # noqa: BLE001
            logging.exception("Failed processing %s", path.name)
            append_failure_log(self.settings.watch_folder, path, str(exc))
            if path.exists() and path.parent == staging_dir(self.settings):
                moved = move_file(path, failed_dir(self.settings))
                logging.error("Moved failed file to %s", moved)
            return {
                "original_filename": original_filename or path.name,
                "assigned_candidate_id": candidate_id,
                "status": "failed",
                "error": str(exc),
            }

    def process(
        self,
        path: Path,
        candidate_id: str | None = None,
        original_filename: str | None = None,
        intake_timestamp: str | None = None,
        batch_id: str | None = None,
    ) -> dict[str, Any]:
        if path.parent != staging_dir(self.settings):
            return {
                "original_filename": original_filename or path.name,
                "assigned_candidate_id": candidate_id,
                "status": "skipped",
                "reason": "not in staging",
            }
        if should_ignore(path):
            logging.debug("Ignoring %s", path.name)
            return {
                "original_filename": original_filename or path.name,
                "assigned_candidate_id": candidate_id,
                "status": "ignored",
                "reason": "unsupported or temporary file",
            }

        logging.info("Detected supported document: %s", path.name)
        if not path.exists():
            logging.info("Ignoring stale filesystem event for already-moved file: %s", path.name)
            return {
                "original_filename": original_filename or path.name,
                "assigned_candidate_id": candidate_id,
                "status": "skipped",
                "reason": "file disappeared",
            }
        wait_until_stable(path, self.settings.stable_checks, self.settings.stable_interval_seconds)

        image_pdf_fallback = False
        try:
            resume_text = extract_resume_text(path)
        except ValueError as exc:
            if str(exc) == NO_EXTRACTED_TEXT and path.suffix.lower() == ".pdf":
                logging.info(
                    "No extractable PDF text found for %s; using rendered-page vision fallback",
                    path.name,
                )
                resume_text = ""
                image_pdf_fallback = True
            else:
                raise

        content_hash = content_hash_for_file(path, None if image_pdf_fallback else resume_text)
        self.processed_log = load_processed_log(self.processed_log_path)
        existing_entry = self.processed_log.get("files", {}).get(content_hash)
        lookup_match = find_lookup_by_hash(self.lookup_worksheet, content_hash)
        if (existing_entry or lookup_match) and not self.settings.allow_reprocess:
            existing_id = (lookup_match or existing_entry or {}).get("candidate_id", "unknown")
            logging.warning(
                "Duplicate resume detected; existing Candidate ID: %s. No new row created.",
                existing_id,
            )
            moved = move_file(path, duplicates_dir(self.settings))
            logging.info("Moved duplicate resume to %s", moved)
            return {
                "original_filename": original_filename or path.name,
                "assigned_candidate_id": candidate_id,
                "existing_candidate_id": existing_id,
                "status": "duplicate",
                "processed_path": str(moved),
            }

        if image_pdf_fallback:
            logging.info("Sending rendered PDF pages to LLM for three-pass structured scoring")
            passes = score_pdf_resume_three_passes_from_images(
                self.client,
                self.settings.openai_model,
                path,
            )
            candidate_name = passes[0]["candidate_name"]
            candidate_email = valid_email(passes[0].get("candidate_email", ""))
        else:
            logging.info("Extracted %s characters from %s", len(resume_text), path.name)
            detection = detect_resume(self.settings, self.client, path, resume_text)
            logging.info(
                "Resume detection for %s: is_resume=%s confidence=%s method=%s reasons=%s",
                path.name,
                detection.is_resume,
                detection.confidence,
                detection.method,
                "; ".join(detection.reasons),
            )
            if not detection.is_resume:
                append_ignored_log(self.settings.watch_folder, path, detection)
                moved = move_file(path, ignored_dir(self.settings))
                logging.info("Moved non-resume document to %s", moved)
                return {
                    "original_filename": original_filename or path.name,
                    "assigned_candidate_id": candidate_id,
                    "status": "ignored_non_resume",
                    "processed_path": str(moved),
                    "detection": {
                        "confidence": detection.confidence,
                        "method": detection.method,
                        "reasons": detection.reasons,
                    },
                }

            logging.info("Sending resume text to LLM for three-pass structured scoring")
            passes = score_resume_three_passes(
                self.client,
                self.settings.openai_model,
                resume_text,
                self.settings.max_resume_chars,
            )
            candidate_name = passes[0]["candidate_name"]
            candidate_email = extract_email(resume_text) or valid_email(passes[0].get("candidate_email", ""))

        if lookup_match and self.settings.allow_reprocess and not self.settings.force_new_candidate_id:
            candidate_id = lookup_match["candidate_id"]
            lookup_row = lookup_match["row"]
            scoring_row = find_scoring_row_by_candidate_id(self.worksheet, candidate_id)
            logging.info("ALLOW_REPROCESS=true; updating existing Candidate ID %s", candidate_id)
        else:
            candidate_id = candidate_id or next_candidate_id(self.lookup_worksheet)
            lookup_row = None
            scoring_row = None
        stored_filename = f"{candidate_id}{path.suffix.lower()}"
        summary = summarize_passes(passes)
        drive_link = None
        if self.drive_service and self.settings.drive_upload_folder_id:
            uploaded = upload_resume_copy(
                self.drive_service,
                path,
                self.settings.drive_upload_folder_id,
                stored_filename,
            )
            drive_link = uploaded.get("webViewLink")

        append_v3_candidate(
            self.worksheet,
            self.lookup_worksheet,
            candidate_id,
            candidate_name,
            candidate_email,
            original_filename or path.name,
            stored_filename,
            drive_link or "",
            content_hash,
            intake_timestamp or "",
            passes,
            summary,
            scoring_row=scoring_row,
            lookup_row=lookup_row,
        )
        selection_summaries = rebuild_selection_board(self.selection_worksheet, self.worksheet)
        result_summary = next(
            (item for item in selection_summaries if item["candidate_id"] == candidate_id),
            None,
        )
        if result_summary:
            logging.info(
                "LATEST RESULT | Candidate ID: %s | Rank: %s | Average Score: %.1f | Status: %s | Confidence: %s | Spread: %s",
                candidate_id,
                result_summary["rank"],
                result_summary["overall_average"],
                result_summary["status"],
                summary["confidence"],
                summary["spread"],
            )
        else:
            logging.warning("Could not find %s on the rebuilt Selection Board", candidate_id)
        moved = move_file(path, processed_dir(self.settings))
        renamed_path = moved.with_name(stored_filename)
        if not renamed_path.exists():
            moved.rename(renamed_path)
            moved = renamed_path
        self.processed_log.setdefault("files", {})[content_hash] = {
            "candidate_id": candidate_id,
            "original_name": original_filename or path.name,
            "original_filename": original_filename or path.name,
            "stored_filename": stored_filename,
            "processed_path": str(moved),
            "candidate_name": candidate_name,
            "batch_id": batch_id,
            "intake_timestamp": intake_timestamp,
            "processed_at": datetime.now(timezone.utc).isoformat(),
        }
        save_processed_log(self.processed_log_path, self.processed_log)
        logging.info("Processed successfully: %s -> %s", path.name, moved)
        return {
            "original_filename": original_filename or path.name,
            "assigned_candidate_id": candidate_id,
            "stored_filename": stored_filename,
            "candidate_name": candidate_name,
            "status": "processed",
            "processed_path": str(moved),
            "overall_average": summary["overall_average"],
            "confidence": summary["confidence"],
            "spread": summary["spread"],
        }


class IntakeHandler(FileSystemEventHandler):
    def __init__(self, settings: Settings):
        self.settings = settings
        self.last_intake_change: datetime | None = None

    def on_created(self, event) -> None:  # noqa: ANN001
        if event.is_directory:
            return
        self.record(Path(event.src_path))

    def on_moved(self, event) -> None:  # noqa: ANN001
        if event.is_directory:
            return
        self.record(Path(event.dest_path))

    def record(self, path: Path) -> None:
        if path.parent != intake_dir(self.settings) or should_ignore(path):
            return
        try:
            if is_convertible_doc_file(path):
                path = convert_doc_to_docx(self.settings, path)
            wait_until_stable(path, self.settings.stable_checks, self.settings.stable_interval_seconds)
            record_intake_file(self.settings, path)
            self.last_intake_change = datetime.now(timezone.utc)
        except Exception as exc:  # noqa: BLE001
            logging.exception("Failed recording intake metadata for %s: %s", path.name, exc)


def record_existing_intake_files(settings: Settings) -> datetime | None:
    latest: datetime | None = None
    for path in intake_dir(settings).iterdir():
        if path.is_file() and not should_ignore(path):
            if is_convertible_doc_file(path):
                path = convert_doc_to_docx(settings, path)
            record_intake_file(settings, path)
            modified = datetime.fromtimestamp(path.stat().st_mtime, timezone.utc)
            latest = max(latest, modified) if latest else modified
    return latest


def eligible_intake_files(settings: Settings) -> list[Path]:
    return [
        path
        for path in intake_dir(settings).iterdir()
        if path.is_file() and is_supported_resume_file(path) and not should_ignore(path)
    ]


def root_drop_files(settings: Settings) -> list[Path]:
    return [
        path
        for path in settings.watch_folder.iterdir()
        if path.is_file() and not should_ignore(path)
    ]


def move_root_drops_to_intake(settings: Settings) -> list[Path]:
    moved_paths: list[Path] = []
    for path in sorted(root_drop_files(settings), key=lambda item: item.name.lower()):
        wait_until_stable(path, settings.stable_checks, settings.stable_interval_seconds)
        if is_convertible_doc_file(path):
            destination = convert_doc_to_docx(settings, path)
        else:
            destination = unique_destination(intake_dir(settings), path.name)
            shutil.move(str(path), str(destination))
            record_intake_file(settings, destination)
        moved_paths.append(destination)
        logging.info("Moved top-level drop into intake: %s -> %s", path.name, destination.name)
    return moved_paths


def move_intake_batch_to_staging(settings: Settings, files: list[Path]) -> list[dict[str, Any]]:
    staged: list[dict[str, Any]] = []
    for path in files:
        wait_until_stable(path, settings.stable_checks, settings.stable_interval_seconds)
        original_filename = path.name
        intake_timestamp = intake_timestamp_for(settings, path)
        staged_path = move_file(path, staging_dir(settings))
        remove_intake_log_entry(settings, original_filename)
        staged.append(
            {
                "original_filename": original_filename,
                "intake_timestamp": intake_timestamp,
                "staged_path": staged_path,
            }
        )
    return staged


def process_batch(settings: Settings) -> dict[str, Any] | None:
    ensure_batch_directories(settings)
    move_root_drops_to_intake(settings)
    record_existing_intake_files(settings)
    files = eligible_intake_files(settings)
    if not files:
        logging.info("No eligible files found in %s", intake_dir(settings))
        return None

    batch_id = next_batch_id(settings)
    seed = random.SystemRandom().randrange(100_000, 1_000_000_000)
    staged = move_intake_batch_to_staging(settings, sorted(files, key=lambda item: item.name.lower()))
    rng = random.Random(seed)
    rng.shuffle(staged)

    handler = ResumeHandler(settings)
    first_id_number = candidate_id_number(next_candidate_id(handler.lookup_worksheet))
    processing_started_at = datetime.now(timezone.utc).isoformat()
    manifest = {
        "batch_id": batch_id,
        "randomization_seed": seed,
        "processing_started_at": processing_started_at,
        "processing_finished_at": None,
        "items": [],
    }

    for offset, item in enumerate(staged):
        candidate_id = candidate_id_from_number(first_id_number + offset)
        staged_path = item["staged_path"]
        candidate_path = staged_path.with_name(f"{candidate_id}{staged_path.suffix.lower()}")
        if candidate_path.exists():
            raise FileExistsError(f"Staging candidate filename already exists: {candidate_path}")
        staged_path.rename(candidate_path)
        item["candidate_id"] = candidate_id
        item["candidate_path"] = candidate_path
        manifest["items"].append(
            {
                "batch_id": batch_id,
                "randomization_seed": seed,
                "original_filename": item["original_filename"],
                "intake_timestamp": item["intake_timestamp"],
                "assigned_candidate_id": candidate_id,
                "staged_filename": candidate_path.name,
                "processing_timestamp": None,
                "status": "queued",
            }
        )

    manifest_path = manifest_dir(settings) / f"{batch_id}.json"
    save_json_file(manifest_path, manifest)
    logging.info(
        "Created randomized batch %s with %s files. Seed=%s Manifest=%s",
        batch_id,
        len(staged),
        seed,
        manifest_path,
    )

    for index, item in enumerate(staged):
        manifest["items"][index]["processing_timestamp"] = datetime.now(timezone.utc).isoformat()
        result = handler.try_process(
            item["candidate_path"],
            candidate_id=item["candidate_id"],
            original_filename=item["original_filename"],
            intake_timestamp=item["intake_timestamp"],
            batch_id=batch_id,
        )
        manifest["items"][index].update(result)
        save_json_file(manifest_path, manifest)

    manifest["processing_finished_at"] = datetime.now(timezone.utc).isoformat()
    save_json_file(manifest_path, manifest)
    logging.info("Finished batch %s. Manifest saved to %s", batch_id, manifest_path)
    return manifest


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="PMP resume intake watcher and batch processor")
    parser.add_argument(
        "--process-batch",
        action="store_true",
        help="Process the current intake folder as one randomized batch, then exit.",
    )
    return parser.parse_args()


def main() -> int:
    setup_logging()
    args = parse_args()
    try:
        settings = load_settings()
        ensure_batch_directories(settings)

        if args.process_batch:
            process_batch(settings)
            return 0

        intake_handler = IntakeHandler(settings)
        intake_handler.last_intake_change = record_existing_intake_files(settings)
        observer = Observer()
        observer.schedule(intake_handler, str(intake_dir(settings)), recursive=False)
        observer.start()
        logging.info("Watching %s for PMP resume intake", intake_dir(settings))
        logging.info(
            "Files will be scored only by --process-batch or after %.1f idle minutes",
            settings.batch_idle_minutes,
        )
        logging.info("Press Ctrl+C to stop")

        while True:
            moved_root_drops = move_root_drops_to_intake(settings)
            if moved_root_drops:
                intake_handler.last_intake_change = datetime.now(timezone.utc)
            files = eligible_intake_files(settings)
            if files:
                if intake_handler.last_intake_change is None:
                    intake_handler.last_intake_change = record_existing_intake_files(settings)
                idle_since = intake_handler.last_intake_change
                idle_seconds = (
                    datetime.now(timezone.utc) - idle_since
                ).total_seconds() if idle_since else 0
                if idle_seconds >= settings.batch_idle_minutes * 60:
                    logging.info(
                        "No new intake files for %.1f minutes; processing randomized batch",
                        settings.batch_idle_minutes,
                    )
                    process_batch(settings)
                    intake_handler.last_intake_change = None
            time.sleep(1)
    except KeyboardInterrupt:
        logging.info("Stopping watcher")
        return 0
    except Exception as exc:  # noqa: BLE001
        logging.exception("Watcher failed: %s", exc)
        return 1
    finally:
        observer = locals().get("observer")
        if observer:
            observer.stop()
            observer.join()


if __name__ == "__main__":
    sys.exit(main())
